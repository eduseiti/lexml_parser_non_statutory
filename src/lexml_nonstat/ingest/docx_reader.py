"""DOCX → :class:`StyledDoc`.

Reads the OOXML directly rather than leaning on ``python-docx``'s convenience
accessors. Three reasons, each forced by the samples:

1. ``Document.paragraphs`` and ``Document.tables`` are separate flat lists, so
   the interleaving is lost — and every sample with a table has it mid-document.
2. Paragraph properties must be resolved through the ``pStyle`` → ``basedOn``
   chain to ``docDefaults``; ``python-docx`` exposes the direct values only.
3. Struck runs and soft breaks need run-level surgery before text is assembled.

Whitespace handling mirrors the reference parser's ``DOCXReader.breakText``
(``lexml-parser-projeto-lei``, Scala), so that Cycle 6b's round-trip has a
common baseline. The one deliberate deviation is NFC normalisation, which the
reference does not perform — see :func:`normalize_text`.
"""

from __future__ import annotations

import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.styles import BabelFish

from .styled import (
    Inline,
    StyledCell,
    StyledDoc,
    StyledPara,
    StyledRow,
    StyledTable,
)

__all__ = [
    "DocxReadError",
    "StyleResolver",
    "normalize_text",
    "read_docx",
]

#: Everything the reference parser treats as horizontal space, including NBSP.
_SPACE_RE = re.compile(r"[ \s\n\r]")
_RUNS_RE = re.compile(r"\s\s+")

#: Word toggle properties are on unless explicitly switched off.
_TOGGLE_OFF = frozenset({"false", "0", "off"})


class DocxReadError(Exception):
    """Raised when a file cannot be read as a DOCX document."""


def normalize_text(text: str) -> str:
    """Normalise to NFC and collapse whitespace.

    The collapse reproduces ``breakText``'s two substitutions::

        text.replaceAll("[\\u00A0\\s\\n\\r]", " ").replaceAll("\\s\\s+", " ")

    so NBSP, tabs and newlines all become ordinary spaces and any run of them
    becomes one. Our pattern omits the explicit ``\\u00A0``: Java's ``\\s`` is
    ASCII-only, which is why the reference must name NBSP separately, whereas
    Python's ``\\s`` already matches it. Same behaviour, one fewer alternative. Leading and trailing spaces survive — the reference emits them
    as explicit ``Space`` segments, and dropping them here would silently join
    words across run boundaries.

    NFC is our addition. The reference works on text that is already composed;
    we cannot assume that of a 300+ document corpus, and a single decomposed
    ``ç`` (``c`` + U+0327) breaks every profile regex, the conservation
    invariant and byte-stable goldens at once — silently, since the two forms
    render identically. All 15 current samples are already NFC, so this is a
    guard for what the corpus will bring, not a fix for what it has.
    """
    return _RUNS_RE.sub(" ", _SPACE_RE.sub(" ", unicodedata.normalize("NFC", text)))


def _toggle_on(element) -> bool:
    """Read an OOXML toggle property (``<w:b/>``, ``<w:strike/>``, …).

    Present with no ``w:val`` means on; ``w:val="false"|"0"|"off"`` means off.
    """
    if element is None:
        return False
    return (element.get(qn("w:val")) or "true").lower() not in _TOGGLE_OFF


def _int_or_none(value: str | None) -> int | None:
    """Word writes measurements as integers, but tolerate junk rather than
    crashing on a malformed document (plan §9.1 robustness layer)."""
    if value is None:
        return None
    try:
        return int(value)
    except ValueError:
        return None


class StyleResolver:
    """Resolves paragraph properties through the style hierarchy.

    Word layers properties: a paragraph's own ``pPr``, then its ``pStyle``, then
    that style's ``basedOn`` ancestors, then ``docDefaults``. Only the first
    layer is directly visible on the paragraph, so anything asking "how far is
    this paragraph indented?" must walk the rest.

    ``basedOn`` chains are followed with a ``seen`` set: a cyclic style graph is
    malformed but not unheard of, and it must terminate rather than hang.
    """

    def __init__(self, document) -> None:
        self._styles = {
            s.get(qn("w:styleId")): s
            for s in document.styles.element
            if s.tag == qn("w:style")
        }
        self._default_para_style = self._find_default_para_style()
        self._doc_default_pPr = self._find_doc_default_pPr(document)
        self._indent_cache: dict[str | None, int | None] = {}

    def _find_default_para_style(self) -> str | None:
        for style_id, element in self._styles.items():
            default = element.get(qn("w:default"))
            if (
                element.get(qn("w:type")) == "paragraph"
                and default is not None
                and default.lower() not in _TOGGLE_OFF
            ):
                return style_id
        return "Normal" if "Normal" in self._styles else None

    @staticmethod
    def _find_doc_default_pPr(document):
        doc_defaults = document.styles.element.find(qn("w:docDefaults"))
        if doc_defaults is None:
            return None
        para_default = doc_defaults.find(qn("w:pPrDefault"))
        return None if para_default is None else para_default.find(qn("w:pPr"))

    def _chain(self, style_id: str | None):
        """Yield the style element chain from ``style_id`` up through
        ``basedOn``, stopping on a cycle."""
        seen: set[str] = set()
        current = style_id
        while current is not None and current not in seen:
            seen.add(current)
            element = self._styles.get(current)
            if element is None:
                return
            yield element
            based_on = element.find(qn("w:basedOn"))
            current = None if based_on is None else based_on.get(qn("w:val"))

    def name(self, style_id: str | None) -> str | None:
        """The human-readable style name (``"Heading 1"``) for a style id.

        Word stores built-in styles under their internal names — CARNE_LEAO's
        headings are ``w:name="heading 1"``, lowercase. ``python-docx`` maps
        those to the UI names users and documentation actually use, and we
        apply the same mapping so ``"Heading 1"`` means the same thing here as
        everywhere else in the ecosystem.
        """
        for element in self._chain(style_id):
            name = element.find(qn("w:name"))
            if name is not None:
                raw = name.get(qn("w:val"))
                return BabelFish.internal2ui(raw) if raw is not None else None
            break  # only the style itself names itself; basedOn does not inherit it
        return None

    def indent(self, style_id: str | None) -> int | None:
        """Left indent in twips, resolved through ``basedOn`` and
        ``docDefaults``. ``None`` when nothing in the chain declares one."""
        key = style_id if style_id is not None else self._default_para_style
        if key in self._indent_cache:
            return self._indent_cache[key]
        result = self._resolve_indent(key)
        self._indent_cache[key] = result
        return result

    def _resolve_indent(self, style_id: str | None) -> int | None:
        for element in self._chain(style_id):
            pPr = element.find(qn("w:pPr"))
            value = _left_indent(pPr)
            if value is not None:
                return value
        return _left_indent(self._doc_default_pPr)

    def outline_level(self, style_id: str | None) -> int | None:
        """``w:outlineLvl`` — 0 for ``Heading 1``, 1 for ``Heading 2``, …

        This is the style's own declaration of depth, which is more reliable
        than parsing the style *name*: a document may rename ``Heading 2`` and
        still carry the correct outline level.
        """
        for element in self._chain(style_id):
            pPr = element.find(qn("w:pPr"))
            if pPr is None:
                continue
            outline = pPr.find(qn("w:outlineLvl"))
            if outline is not None:
                return _int_or_none(outline.get(qn("w:val")))
        return None

    @property
    def default_style_id(self) -> str | None:
        return self._default_para_style


def _left_indent(pPr) -> int | None:
    """``w:ind/@w:left`` (or its ``@w:start`` alias in newer Word)."""
    if pPr is None:
        return None
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return None
    return _int_or_none(ind.get(qn("w:left")) or ind.get(qn("w:start")))


def _strip_struck_runs(paragraph) -> None:
    """Remove every ``<w:r>`` carrying an active strike, in place.

    Must run *before* soft-break splitting. The reference parser's comment
    explains why: a single struck run can contain ``<w:br/>``-separated text, so
    splitting first would leave the post-break segments outside the run's
    ``<w:rPr>`` and let struck text leak into the output.

    In ``sumula_stj_125`` this drops the ordinal ``ª`` in citations like
    ``(2ª T, 03.08.1994)``, matching the reference parser's rendering.
    """
    for run in list(paragraph.findall(qn("w:r"))):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            continue
        if _toggle_on(rPr.find(qn("w:strike"))) or _toggle_on(rPr.find(qn("w:dstrike"))):
            paragraph.remove(run)


def _run_inlines(run, href: str | None) -> list[Inline]:
    """Text of one run as inlines, split at soft breaks.

    Returns one inline per ``<w:br/>``-delimited segment; the caller uses the
    segment boundaries to split the paragraph. A ``<w:tab/>`` becomes a space,
    which ``normalize_text`` then collapses.
    """
    rPr = run.find(qn("w:rPr"))
    bold = italic = sup = sub = False
    if rPr is not None:
        bold = _toggle_on(rPr.find(qn("w:b")))
        italic = _toggle_on(rPr.find(qn("w:i")))
        vert_align = rPr.find(qn("w:vertAlign"))
        if vert_align is not None:
            align = vert_align.get(qn("w:val"))
            sup = align == "superscript"
            sub = align == "subscript"

    segments: list[str] = [""]
    for child in run:
        tag = child.tag
        if tag in (qn("w:t"), qn("w:delText")):
            segments[-1] += child.text or ""
        elif tag == qn("w:tab"):
            segments[-1] += " "
        elif tag == qn("w:br"):
            # A page/column break is not a line break within the paragraph.
            if (child.get(qn("w:type")) or "textWrapping") == "textWrapping":
                segments.append("")
            else:
                segments[-1] += " "
        elif tag == qn("w:cr"):
            segments.append("")
        elif tag == qn("w:noBreakHyphen"):
            segments[-1] += "-"
        elif tag == qn("w:softHyphen"):
            segments[-1] += ""
        elif tag == qn("w:sym"):
            char = child.get(qn("w:char"))
            if char:
                try:
                    segments[-1] += chr(int(char, 16))
                except ValueError:
                    pass

    return [
        Inline(text=s, bold=bold, italic=italic, sup=sup, sub=sub, href=href)
        for s in segments
    ]


def _paragraph_segments(paragraph, rels) -> list[list[Inline]]:
    """The paragraph's inlines, split into one list per soft-break-delimited
    line. A paragraph with no soft break yields exactly one list."""
    lines: list[list[Inline]] = [[]]

    def consume(run, href: str | None) -> None:
        parts = _run_inlines(run, href)
        for position, inline in enumerate(parts):
            if position:
                lines.append([])
            if inline.text:
                lines[-1].append(inline)

    for child in paragraph:
        if child.tag == qn("w:r"):
            consume(child, None)
        elif child.tag == qn("w:hyperlink"):
            href = _hyperlink_target(child, rels)
            for run in child.findall(qn("w:r")):
                consume(run, href)
        elif child.tag in (qn("w:ins"), qn("w:smartTag"), qn("w:sdt")):
            # Tracked insertions and wrappers: descend, they hold ordinary runs.
            for run in child.iter(qn("w:r")):
                consume(run, None)

    return lines


def _hyperlink_target(hyperlink, rels) -> str | None:
    """Resolve ``<w:hyperlink r:id>`` through the part's relationships."""
    rel_id = hyperlink.get(qn("r:id"))
    if rel_id and rels is not None and rel_id in rels:
        return rels[rel_id].target_ref
    anchor = hyperlink.get(qn("w:anchor"))
    return f"#{anchor}" if anchor else None


def _merge_inlines(inlines: list[Inline]) -> tuple[Inline, ...]:
    """Normalise text and merge adjacent runs sharing formatting.

    Word splits runs on spell-check state, language, revision marks and other
    invisible boundaries, so a single sentence routinely arrives as a dozen
    identically-formatted runs. Merging keeps goldens legible and makes
    ``Inline`` boundaries meaningful — a boundary now marks a real formatting
    change, which is exactly what Cycle 4's typography evidence reads.
    """
    merged: list[Inline] = []
    for inline in inlines:
        text = normalize_text(inline.text)
        if not text:
            continue
        candidate = Inline(
            text=text,
            bold=inline.bold,
            italic=inline.italic,
            sup=inline.sup,
            sub=inline.sub,
            href=inline.href,
        )
        if merged:
            last = merged[-1]
            same_format = (
                last.bold == candidate.bold
                and last.italic == candidate.italic
                and last.sup == candidate.sup
                and last.sub == candidate.sub
                and last.href == candidate.href
            )
            if same_format:
                joined = normalize_text(last.text + candidate.text)
                merged[-1] = Inline(
                    text=joined,
                    bold=last.bold,
                    italic=last.italic,
                    sup=last.sup,
                    sub=last.sub,
                    href=last.href,
                )
                continue
        merged.append(candidate)

    if merged:
        first, last = merged[0], merged[-1]
        merged[0] = Inline(
            text=first.text.lstrip(" "),
            bold=first.bold,
            italic=first.italic,
            sup=first.sup,
            sub=first.sub,
            href=first.href,
        )
        last = merged[-1]
        merged[-1] = Inline(
            text=last.text.rstrip(" "),
            bold=last.bold,
            italic=last.italic,
            sup=last.sup,
            sub=last.sub,
            href=last.href,
        )
    return tuple(i for i in merged if i.text)


def _read_paragraph(
    paragraph, resolver: StyleResolver, rels, index: int, *, drop_strikethrough: bool
) -> list[StyledPara]:
    """One OOXML paragraph as one or more :class:`StyledPara`.

    More than one when the paragraph contains soft breaks: Word treats
    ``<w:br/>`` as a line break inside a paragraph, but for our purposes each
    line is a separate block — the reference parser splits the same way, and
    leaving them joined would merge a heading with the text beneath it.
    """
    if drop_strikethrough:
        _strip_struck_runs(paragraph)

    pPr = paragraph.find(qn("w:pPr"))
    style_id = None
    num_id = ilvl = None
    indent_direct = first_line = hanging = None
    alignment = None

    if pPr is not None:
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            style_id = pStyle.get(qn("w:val"))

        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            num_element = numPr.find(qn("w:numId"))
            ilvl_element = numPr.find(qn("w:ilvl"))
            if num_element is not None:
                num_id = num_element.get(qn("w:val"))
            if ilvl_element is not None:
                ilvl = _int_or_none(ilvl_element.get(qn("w:val")))

        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            indent_direct = _int_or_none(
                ind.get(qn("w:left")) or ind.get(qn("w:start"))
            )
            first_line = _int_or_none(ind.get(qn("w:firstLine")))
            hanging = _int_or_none(ind.get(qn("w:hanging")))

        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            alignment = jc.get(qn("w:val"))

    inherited = resolver.indent(style_id)
    indent_effective = (
        indent_direct if indent_direct is not None else (inherited or 0)
    )

    common = {
        "style": resolver.name(style_id) if style_id else resolver.name(
            resolver.default_style_id
        ),
        "style_id": style_id,
        "outline_level": resolver.outline_level(style_id),
        "num_id": num_id,
        "ilvl": ilvl,
        "indent_direct": indent_direct,
        "indent_effective": indent_effective,
        "first_line": first_line,
        "hanging": hanging,
        "alignment": alignment,
    }

    lines = _paragraph_segments(paragraph, rels)
    result: list[StyledPara] = []
    for line in lines:
        result.append(
            StyledPara(inlines=_merge_inlines(line), index=index + len(result), **common)
        )
    return result or [StyledPara(index=index, **common)]


def _read_table(
    table, resolver: StyleResolver, rels, index: int, *, drop_strikethrough: bool
) -> StyledTable:
    rows: list[StyledRow] = []
    for tr in table.findall(qn("w:tr")):
        cells: list[StyledCell] = []
        for tc in tr.findall(qn("w:tc")):
            paras: list[StyledPara] = []
            for p in tc.findall(qn("w:p")):
                paras.extend(
                    _read_paragraph(
                        p,
                        resolver,
                        rels,
                        len(paras),
                        drop_strikethrough=drop_strikethrough,
                    )
                )
            cells.append(StyledCell(paras=tuple(paras)))
        rows.append(StyledRow(cells=tuple(cells)))
    return StyledTable(rows=tuple(rows), index=index)


def read_docx(
    path: str | Path, *, drop_strikethrough: bool = True
) -> StyledDoc:
    """Read a ``.docx`` file into a :class:`StyledDoc`.

    Args:
        path: the document to read.
        drop_strikethrough: discard struck-through runs, as the reference
            parser does. Set ``False`` to retain them — useful when reviewing
            what a document actually contains, since struck text is invisible
            in the default output.

    Raises:
        DocxReadError: the file is missing or is not a readable DOCX.
    """
    path = Path(path)
    if not path.exists():
        raise DocxReadError(f"no such file: {path}")
    try:
        document = Document(str(path))
    except Exception as exc:  # python-docx raises a variety of types
        raise DocxReadError(f"cannot read {path.name} as DOCX: {exc}") from exc

    resolver = StyleResolver(document)
    rels = document.part.rels
    blocks: list[StyledPara | StyledTable] = []

    for child in document.element.body:
        if child.tag == qn("w:p"):
            blocks.extend(
                _read_paragraph(
                    child,
                    resolver,
                    rels,
                    len(blocks),
                    drop_strikethrough=drop_strikethrough,
                )
            )
        elif child.tag == qn("w:tbl"):
            blocks.append(
                _read_table(
                    child,
                    resolver,
                    rels,
                    len(blocks),
                    drop_strikethrough=drop_strikethrough,
                )
            )

    # Indices are assigned after soft-break splitting, so they address the
    # blocks that actually exist rather than the source paragraphs.
    renumbered: list[StyledPara | StyledTable] = []
    for position, block in enumerate(blocks):
        if isinstance(block, StyledTable):
            renumbered.append(StyledTable(rows=block.rows, index=position))
        else:
            renumbered.append(
                StyledPara(
                    inlines=block.inlines,
                    style=block.style,
                    style_id=block.style_id,
                    outline_level=block.outline_level,
                    num_id=block.num_id,
                    ilvl=block.ilvl,
                    indent_direct=block.indent_direct,
                    indent_effective=block.indent_effective,
                    first_line=block.first_line,
                    hanging=block.hanging,
                    alignment=block.alignment,
                    index=position,
                )
            )

    return StyledDoc(blocks=tuple(renumbered), source=path.name)
