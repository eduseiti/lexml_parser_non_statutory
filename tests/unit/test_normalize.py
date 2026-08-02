"""`normalize_text` invariants: one Unicode form, one kind of space.

Two things must hold for every string that enters the pipeline, and both are
invisible when they break:

1. **NFC.** ``"ç"`` written as ``c`` + U+0327 renders identically to U+00E7 but
   compares unequal, so a decomposed cedilla silently defeats every profile
   regex, the text-conservation invariant and byte-stable goldens at once. The
   corpus is already NFC today; the normaliser is a guard for what arrives
   later. `test_nfc_is_not_a_noop` exists to fail loudly if that guard is ever
   removed, and `test_all_samples_are_nfc` is the tripwire on the corpus.

2. **breakText parity.** Whitespace collapsing must reproduce the Scala
   reference parser's ``DOCXReader.breakText``::

       text.replaceAll("[ \\s\\n\\r]", " ").replaceAll("\\s\\s+", " ")

   Cycle 6b's round-trip compares our output against that parser's, so any
   divergence here becomes a diff there. Note what the rules do *not* do:
   leading and trailing spaces are preserved, because the reference emits them
   as explicit ``Space`` segments and dropping them would join words across run
   boundaries.
"""

from __future__ import annotations

import unicodedata
from pathlib import Path

import pytest
from docx import Document

from lexml_nonstat.ingest import normalize_text

from tests.conftest import REPO_ROOT

NBSP = " "

SAMPLES = sorted((REPO_ROOT / "samples").glob("*.docx"))


# --------------------------------------------------------------------------
# NFC composition
# --------------------------------------------------------------------------


def test_nfc_composes_cedilla():
    """`c` + combining cedilla is the failure mode we most expect: "Seção"."""
    result = normalize_text("ç")

    assert result == "ç"
    assert len(result) == 1


def test_nfc_composes_tilde():
    """"ão" — the single most common decomposable sequence in the corpus."""
    result = normalize_text("ão")

    assert result == "ão"
    assert len(result) == 2


def test_nfc_composes_acute():
    """Uppercase too: "União" opens with a composable `U` + acute in some
    documents produced by non-Word editors."""
    result = normalize_text("Ú")

    assert result == "Ú"
    assert len(result) == 1


def test_nfc_is_not_a_noop():
    """The normaliser must actually do something.

    This test's entire purpose is to fail if the
    ``unicodedata.normalize("NFC", ...)`` call is deleted from
    ``normalize_text``. The two forms render identically, so assertions are on
    code points and length, never on visual equality: `raw` and `expected`
    below are *equal to the eye* and unequal to `==`.
    """
    composed = "Advocacia-Geral da União — Seção de Anulação"
    raw = unicodedata.normalize("NFD", composed)

    # Preconditions: the input really is decomposed, and decomposition is
    # observable (otherwise the test would pass vacuously).
    assert raw != composed
    assert len(raw) > len(composed)
    assert "̃" in raw and "̧" in raw

    result = normalize_text(raw)

    assert result != raw, "NFC normalisation appears to have been removed"
    assert result == composed
    assert len(result) == len(composed)
    assert [ord(c) for c in result] == [ord(c) for c in composed]
    assert all(unicodedata.combining(c) == 0 for c in result)


@pytest.mark.parametrize(
    "text",
    [
        "PARECER n. 93/2018/DECOR/CGU/AGU",
        "SEÇÃO II — DAS OBRIGAÇÕES ACESSÓRIAS",
        "Carnê-Leão",
        "Súmula CARF nº 42",
        "Advocacia-Geral da União",
        "Procuradoria-Geral da Fazenda Nacional",
    ],
)
def test_decomposed_legal_text_is_recomposed(text):
    """Round trip through NFD and back for real strings from the samples."""
    assert normalize_text(unicodedata.normalize("NFD", text)) == text


# --------------------------------------------------------------------------
# Whitespace collapsing — breakText parity
# --------------------------------------------------------------------------


def test_nbsp_becomes_space():
    """NBSP is not a distinguishable character downstream; it is a space."""
    result = normalize_text(f"a{NBSP}b")

    assert result == "a b"
    assert NBSP not in result


def test_runs_of_spaces_collapse():
    assert normalize_text("a    b") == "a b"


def test_newlines_and_tabs_collapse():
    assert normalize_text("a\n\tb") == "a b"


def test_mixed_whitespace_collapses():
    """Heterogeneous runs collapse as one: the second regex sees the output of
    the first, by which point every space-ish character is already `" "`."""
    assert normalize_text("a  \n b") == "a b"


def test_empty_and_whitespace_only():
    """A whitespace-only string collapses to one space, it is *not* stripped.

    The reference parser emits leading/trailing whitespace as explicit `Space`
    segments; stripping here would let words from adjacent runs run together.
    """
    assert normalize_text("") == ""
    assert normalize_text("   ") == " "


def test_internal_single_spaces_are_untouched():
    """The common case must be a pass-through — no word boundary is invented
    or removed in already-clean text."""
    text = "Advocacia-Geral da União manifesta-se pela aprovação do parecer."

    assert normalize_text(text) == text


@pytest.mark.parametrize(
    ("label", "raw", "expected"),
    [
        ("single space kept", "a b", "a b"),
        ("two spaces", "a  b", "a b"),
        ("many spaces", "a" + " " * 12 + "b", "a b"),
        ("tab", "a\tb", "a b"),
        ("consecutive tabs", "a\t\tb", "a b"),
        ("newline", "a\nb", "a b"),
        ("carriage return", "a\rb", "a b"),
        ("crlf", "a\r\nb", "a b"),
        ("vertical tab", "a\x0bb", "a b"),
        ("form feed", "a\x0cb", "a b"),
        ("nbsp", f"a{NBSP}b", "a b"),
        ("run of nbsp", f"a{NBSP}{NBSP}{NBSP}b", "a b"),
        ("nbsp mixed with space", f"a {NBSP} b", "a b"),
        ("leading space preserved", " a", " a"),
        ("trailing space preserved", "a ", "a "),
        ("leading run collapsed not stripped", "   a", " a"),
        ("trailing run collapsed not stripped", "a   ", "a "),
        ("both ends preserved", "  a  ", " a "),
        ("leading newline becomes a space", "\nArt. 1º", " Art. 1º"),
        ("whitespace only", " \t\n ", " "),
        ("nbsp only", NBSP, " "),
        ("no whitespace at all", "PARECER", "PARECER"),
        ("empty", "", ""),
        (
            "realistic run-split text",
            "PARECER\tn. 93/2018\nDECOR/CGU/AGU",
            "PARECER n. 93/2018 DECOR/CGU/AGU",
        ),
    ],
)
def test_matches_breaktext_rules(label, raw, expected):
    """Cases derived directly from the two Scala substitutions.

    ``[ \\s\\n\\r]`` → `" "` maps every space-ish code point (Python's `\\s`
    covers NBSP, tab, CR, LF, VT and FF for `str` patterns, matching the
    reference's explicit character class), then ``\\s\\s+`` → `" "` collapses
    any surviving run. Neither substitution is anchored, so the ends of the
    string are collapsed but never trimmed.
    """
    assert normalize_text(raw) == expected, label


@pytest.mark.parametrize(
    "raw",
    [
        "",
        " ",
        "   ",
        "a b",
        "a  \n\t b",
        f"  Seção{NBSP}{NBSP}II  ",
        "çaõ   \t  Ú",
        "PARECER n. 93/2018/DECOR/CGU/AGU",
    ],
)
def test_normalisation_is_idempotent(raw):
    """Invariant #4 (determinism) at the character level.

    `_merge_inlines` re-normalises concatenations of already-normalised text,
    so a second pass must be a no-op — otherwise merging two inlines could
    change text that was already settled.
    """
    once = normalize_text(raw)

    assert normalize_text(once) == once


def test_no_whitespace_other_than_plain_space_survives():
    """Whatever goes in, only U+0020 comes out."""
    result = normalize_text(f"a\t\n\r\x0b\x0c{NBSP}b")

    assert result == "a b"
    assert not any(c.isspace() and c != " " for c in result)


# --------------------------------------------------------------------------
# Corpus tripwire
# --------------------------------------------------------------------------


def test_sample_corpus_is_complete():
    """The tripwire below is only meaningful if it covers every sample."""
    assert len(SAMPLES) == 15, [p.name for p in SAMPLES]


@pytest.mark.parametrize("sample", SAMPLES, ids=lambda p: p.stem)
def test_all_samples_are_nfc(sample: Path):
    """Every paragraph in every sample is already NFC.

    This documents the state of today's corpus rather than a property of the
    parser: all 15 documents came out of Word, which composes. The value is as
    a tripwire — when the first decomposed document arrives, this test names it,
    and the failure is a note that the corpus changed, not a bug in
    `normalize_text` (which handles it). Parametrised per file so the report
    identifies the offending document.
    """
    document = Document(str(sample))
    texts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)

    offenders = [
        text for text in texts if text != unicodedata.normalize("NFC", text)
    ]

    assert not offenders, (
        f"{sample.name} contains {len(offenders)} non-NFC paragraph(s); "
        f"first: {offenders[0]!r}"
    )
