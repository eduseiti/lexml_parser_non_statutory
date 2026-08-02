"""`read_docx` extracts the structural evidence later cycles classify on.

The invariant guarded here: every signal that Cycles 3 and 4 will read to decide
"is this a heading / a list item / a quotation?" survives ingestion intact and
in source order — style names, outline levels, numbering, both indentation
readings, run formatting, hyperlink targets, tables and their position.

Values asserted below are *measurements* of the 15 real samples, not guesses.
Where a number is surprising (the lone 240-twip article, the 2909 style trap)
the docstring says why, so a future failure can be triaged as "the corpus
changed" versus "the reader regressed".
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from lxml import etree

from lexml_nonstat.ingest import DocxReadError, StyledPara, StyledTable, read_docx
from lexml_nonstat.ingest.docx_reader import StyleResolver, _read_paragraph

# --------------------------------------------------------------------------
# Sample inventory
# --------------------------------------------------------------------------

CARNE_LEAO = "sistema_de_recolhimento_mensal_obrigatorio_CARNE_LEAO.docx"
PARECER_93 = "parecer_93_2018_decor_cgu_agu.docx"
SUMULA_STJ = "sumula_stj_125.docx"
RESP = "REsp_1306393.docx"
PAR_COSIT = "par_cosit_26_20000629.docx"
PN_CST_38 = "pn_cst_38_19801031.docx"

ALL_SAMPLES = (
    RESP,
    "ad_pgfn_13_20111220.docx",
    "ad_pgfn_3_20080918.docx",
    "ad_srf_22_19970430.docx",
    "ad_srf_3_19990107.docx",
    "adn_cosit_19_20001025.docx",
    "adn_cst_10_19910417.docx",
    PAR_COSIT,
    PARECER_93,
    PN_CST_38,
    "port_mf_277_20180607.docx",
    "port_mf_454_19770825.docx",
    CARNE_LEAO,
    "sumula_carf_42.docx",
    SUMULA_STJ,
)

#: (sample, block index of the sole table, (rows, cols)). Every other sample has
#: no table at all — see `test_tables_extracted`.
TABLE_SAMPLES = (
    (RESP, 4, (5, 2)),
    (SUMULA_STJ, 10, (7, 4)),
    (PAR_COSIT, 11, (2, 3)),
)

#: `w:ind/@w:left` band the plan predicts for quoted statute articles.
QUOTE_BAND = range(2880, 2931)

_ARTICLE_RE = re.compile(r"^\s*Art(igo)?\.?\s*\d", re.IGNORECASE)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = f'xmlns:w="{_W}"'


# --------------------------------------------------------------------------
# Fixtures
#
# `read_docx` is not cached and the two large samples (450 and 397 blocks) are
# each read by half a dozen tests. Module-scoped fixtures keep the file under a
# couple of seconds; `StyledDoc` is frozen, so sharing is safe.
# --------------------------------------------------------------------------


@pytest.fixture(scope="session")
def samples_dir(repo_root: Path) -> Path:
    return repo_root / "samples"


@pytest.fixture(scope="module")
def carne_leao(samples_dir: Path):
    return read_docx(samples_dir / CARNE_LEAO)


@pytest.fixture(scope="module")
def parecer_93(samples_dir: Path):
    return read_docx(samples_dir / PARECER_93)


@pytest.fixture(scope="module")
def sumula_stj(samples_dir: Path):
    return read_docx(samples_dir / SUMULA_STJ)


@pytest.fixture(scope="module")
def sumula_stj_struck(samples_dir: Path):
    """`sumula_stj_125` read with struck runs *retained*."""
    return read_docx(samples_dir / SUMULA_STJ, drop_strikethrough=False)


def _synthetic_resolver(styles_xml: str) -> StyleResolver:
    """A `StyleResolver` over an in-memory `styles.xml`.

    `StyleResolver` only ever touches `document.styles.element`, so two stub
    objects are enough — and no synthetic `.docx` has to be written into
    `samples/`, which is reserved for real corpus documents.
    """
    element = etree.fromstring(styles_xml.encode("utf-8"))

    class _Styles:
        def __init__(self, el):
            self.element = el

    class _Document:
        def __init__(self, el):
            self.styles = _Styles(el)

    return StyleResolver(_Document(element))


# --------------------------------------------------------------------------
# Headings and outline levels
# --------------------------------------------------------------------------


def test_carne_leao_headings(carne_leao):
    """CARNE_LEAO's six headings, in document order.

    Word stores built-in styles under their *internal* names (`heading 1`,
    lowercase); the reader maps them to the UI names, so the assertion is on
    `"Heading 1"`.
    """
    headings = [
        p
        for p in carne_leao.paragraphs
        if p.style in ("Heading 1", "Heading 2")
    ]

    assert [p.style for p in headings] == [
        "Heading 1",
        "Heading 2",
        "Heading 2",
        "Heading 2",
        "Heading 2",
        "Heading 2",
    ]
    assert headings[0].text.startswith("Sistema de Recolhimento Mensal Obrigatório")
    assert [p.text for p in headings[1:]] == [
        "O que é?",
        "Quem pode utilizar este serviço?",
        "Etapas para a realização deste serviço",
        "Outras Informações",
        "Lei Geral de Proteção de Dados Pessoais - LGPD",
    ]
    # Document order, not accessor order.
    assert [p.index for p in headings] == sorted(p.index for p in headings)


def test_sumula_stj_heading_styles(sumula_stj):
    """39 `Heading 1` paragraphs — one per cited precedent block."""
    headings = [p for p in sumula_stj.paragraphs if p.style == "Heading 1"]

    assert len(headings) == 39
    assert all(not p.is_empty for p in headings)


@pytest.mark.parametrize("style, expected_level", [("Heading 1", 0), ("Heading 2", 1)])
def test_outline_level_resolved(carne_leao, style, expected_level):
    """`outline_level` comes from the *style's* `w:outlineLvl`, resolved through
    `basedOn` — the paragraphs themselves declare none."""
    matching = [p for p in carne_leao.paragraphs if p.style == style]

    assert matching, f"no {style} paragraph in CARNE_LEAO"
    assert {p.outline_level for p in matching} == {expected_level}
    assert all(p.indent_direct is None or isinstance(p.indent_direct, int)
               for p in matching)


# --------------------------------------------------------------------------
# Numbering
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "sample, expected",
    [(PARECER_93, 197), (CARNE_LEAO, 39), (SUMULA_STJ, 14)],
)
def test_numbering_captured(samples_dir, sample, expected):
    """`w:numPr/w:numId` survives ingestion, and `is_listed` agrees with it."""
    doc = read_docx(samples_dir / sample)
    listed = [p for p in doc.paragraphs if p.num_id is not None]

    assert len(listed) == expected
    assert all(p.is_listed for p in listed)
    assert all(isinstance(p.num_id, str) for p in listed)


@pytest.mark.parametrize(
    "sample, expected_levels",
    [(PARECER_93, {0, 1, 2, 3}), (CARNE_LEAO, {0, 1, 2}), (SUMULA_STJ, {0, 1})],
)
def test_nested_ilvl_distinguished(samples_dir, sample, expected_levels):
    """Nested list levels are distinguishable, not flattened to one.

    `ilvl` is what Cycle 3 reads to nest `Item`s inside `Inciso`s; collapsing
    every level to 0 would be invisible in the text but destroy the hierarchy.
    """
    doc = read_docx(samples_dir / sample)
    levels = {p.ilvl for p in doc.paragraphs if p.ilvl is not None}

    assert levels == expected_levels
    assert len(levels) >= 2
    assert all(isinstance(level, int) for level in levels)


# --------------------------------------------------------------------------
# Indentation — the load-bearing quotation discriminator (spec §2.1)
# --------------------------------------------------------------------------


def test_parecer_93_quoted_articles_indent(parecer_93):
    """Exactly 21 `Art.`-initial paragraphs; 20 sit in the 2880–2930 quote band.

    The 21st is an outlier at `indent_direct == 240`:
    `"Art. 4º - São atribuições do Advogado-Geral da União:"`. It is a *quoted*
    article like the rest, but Word records it with a different direct indent,
    so indentation alone cannot classify all 21 — which is precisely why Cycle 4
    must combine this signal with others rather than thresholding on it.
    """
    articles = [
        p for p in parecer_93.paragraphs if _ARTICLE_RE.match(p.text.strip())
    ]

    assert len(articles) == 21
    in_band = [p for p in articles if p.indent_direct in QUOTE_BAND]
    assert len(in_band) >= 20, [p.indent_direct for p in articles]

    outliers = [p for p in articles if p.indent_direct not in QUOTE_BAND]
    assert [p.indent_direct for p in outliers] == [240]
    assert outliers[0].text.strip().startswith("Art. 4º")
    assert {p.indent_direct for p in articles} <= {240, 2880, 2908, 2930}


def test_parecer_93_indent_bands_are_disjoint(parecer_93):
    """A small-indent population exists, disjoint from the quote band.

    NOTE — spec §2.1 says "the modal indent is not 0" and §6 predicted the modal
    *non-quote* direct indent would be < 100. Measured, it is 2908: that band
    holds 131 paragraphs, not just the 21 `Art.`-initial ones, because the whole
    quoted excerpt (its incisos, alíneas and continuation paragraphs) carries the
    same indent as the articles heading it. So 2908 is the modal direct indent of
    the *document*, and the discriminator is not "quote band vs modal" but
    "quote band vs the small band".

    What is true and load-bearing is asserted here: a distinct small-indent
    population (7–65 twips, front matter and unquoted commentary) exists, and the
    two bands do not overlap — an order of magnitude apart, so a Cycle 4
    threshold anywhere in 100–2800 separates them.
    """
    body = [
        p.indent_direct
        for p in parecer_93.paragraphs
        if not p.is_empty and p.indent_direct is not None
    ]

    assert body, "no body paragraph carries a direct indent"
    small = [v for v in body if v < 100]
    quoted = [v for v in body if v in QUOTE_BAND]

    assert len(small) >= 30, f"small-indent population vanished: {sorted(set(body))}"
    assert max(small) < 100 <= 2800 < min(quoted)
    # The band is broader than the 21 articles — the quoted excerpt shares it.
    assert len(quoted) > 21

    modal = max(set(body), key=body.count)
    assert modal in QUOTE_BAND, (
        f"modal direct indent moved to {modal}; §2.1's correction assumed 2908"
    )


def test_indent_direct_and_effective_differ(parecer_93):
    """Both indent readings are present and neither substitutes for the other.

    219 non-empty paragraphs declare no direct `w:ind` at all; their effective
    indent is inherited (2909 from `Normal`, or 32 from a nested style). 2909 is
    one twip from the quote band's 2908 — read *only* effective indent and the
    quotation signal disappears, hence Q1's decision to keep both.
    """
    inherited = [
        p
        for p in parecer_93.paragraphs
        if p.indent_direct is None and not p.is_empty
    ]

    assert len(inherited) == 219
    assert {p.indent_effective for p in inherited} == {32, 2909}
    assert any(p.indent_effective == 2909 for p in inherited)
    # The trap made explicit: inherited 2909 is adjacent to the 2908 quote band.
    assert 2908 in {
        p.indent_direct for p in parecer_93.paragraphs if p.indent_direct
    }


def test_style_indent_inherited(samples_dir):
    """`Normal` in `parecer_93` declares `w:ind/@w:left="2909"`; the resolver
    reports it for the style id and for `None` (which means "default style")."""
    from docx import Document

    resolver = StyleResolver(Document(str(samples_dir / PARECER_93)))

    assert resolver.indent("Normal") == 2909
    assert resolver.default_style_id == "Normal"
    assert resolver.indent(None) == 2909
    assert resolver.name("Normal") == "Normal"


def test_style_cycle_terminates():
    """A cyclic `basedOn` graph must return, not hang.

    Synthetic on purpose: no real sample has a cycle, and one malformed document
    in the corpus would otherwise wedge the whole pipeline. The chain A→B→A is
    walked with a `seen` set, so resolution falls through to `docDefaults`.
    """
    resolver = _synthetic_resolver(
        f"""<w:styles {_NS}>
          <w:docDefaults><w:pPrDefault><w:pPr>
            <w:ind w:left="99"/>
          </w:pPr></w:pPrDefault></w:docDefaults>
          <w:style w:type="paragraph" w:styleId="A">
            <w:name w:val="A"/><w:basedOn w:val="B"/>
          </w:style>
          <w:style w:type="paragraph" w:styleId="B">
            <w:name w:val="B"/><w:basedOn w:val="A"/>
          </w:style>
        </w:styles>"""
    )

    # Each call must terminate; the assertions also pin the fall-through result.
    assert resolver.indent("A") == 99
    assert resolver.indent("B") == 99
    assert resolver.outline_level("A") is None
    assert resolver.name("A") == "A"
    assert list(resolver._chain("A")) != []


def test_style_indent_falls_back_to_doc_defaults():
    """No sample relies on `docDefaults`-only indent today, so it is exercised
    synthetically (spec §8 risk row)."""
    resolver = _synthetic_resolver(
        f"""<w:styles {_NS}>
          <w:docDefaults><w:pPrDefault><w:pPr>
            <w:ind w:left="1440"/>
          </w:pPr></w:pPrDefault></w:docDefaults>
          <w:style w:type="paragraph" w:styleId="Body">
            <w:name w:val="Body"/>
          </w:style>
        </w:styles>"""
    )

    assert resolver.indent("Body") == 1440
    assert resolver.indent("NoSuchStyle") == 1440


# --------------------------------------------------------------------------
# Run formatting
# --------------------------------------------------------------------------


def test_superscript_runs_preserved(parecer_93):
    """305 superscript runs — the ordinals (`4º`, `1ª`) that distinguish an
    article number from a plain digit."""
    sups = [i for p in parecer_93.paragraphs for i in p.inlines if i.sup]

    assert len(sups) == 305
    assert all(not i.sub for i in sups)
    assert any(i.text.strip() in {"º", "ª", "o", "a"} for i in sups)


def test_bold_italic_preserved(sumula_stj):
    """Both bold and italic runs survive, and `is_plain` reflects them."""
    inlines = [i for p in sumula_stj.paragraphs for i in p.inlines]
    bold = [i for i in inlines if i.bold]
    italic = [i for i in inlines if i.italic]

    assert bold, "no bold run detected in sumula_stj_125"
    assert italic, "no italic run detected in sumula_stj_125"
    assert all(not i.is_plain for i in bold + italic)


# --------------------------------------------------------------------------
# Struck runs (Q2)
# --------------------------------------------------------------------------


def _table_cell_texts(doc) -> list[str]:
    return [c.text for t in doc.tables for row in t.rows for c in row.cells]


def test_struck_runs_dropped(sumula_stj, sumula_stj_struck):
    """The 18 struck runs live *inside the table*, not in body paragraphs.

    They are ordinal markers in court citations: `(2ª T, …)` is struck down to
    `(2 T, …)`. Dropping them matches the reference parser
    (`DOCXReader.stripStruckRuns`); `drop_strikethrough=False` restores them.
    """
    dropped = _table_cell_texts(sumula_stj)
    kept = _table_cell_texts(sumula_stj_struck)

    assert len(dropped) == len(kept)
    differing = [(d, k) for d, k in zip(dropped, kept) if d != k]
    assert differing, "no cell changed — struck runs are not being dropped"

    assert "(2 T, 03.08.1994 — DJ 22.08.1994)" in dropped
    assert "(2ª T, 03.08.1994 — DJ 22.08.1994)" in kept
    assert "(2ª T, 03.08.1994 — DJ 22.08.1994)" not in dropped

    # Every difference is a lost ordinal, never a lost word.
    for struck_out, retained in differing:
        assert len(struck_out) < len(retained)


def test_strike_toggle_off_is_kept():
    """`<w:strike w:val="false"/>` is a toggle switched *off* — the run stays.

    Word writes this whenever a style turns strike on and a run turns it back
    off. Treating any `<w:strike>` element as "struck" would silently delete
    live text, so the toggle values `false`/`0`/`off` must all read as on-page.
    """
    resolver = _synthetic_resolver(f"<w:styles {_NS}></w:styles>")
    paragraph = etree.fromstring(
        f"""<w:p {_NS}>
          <w:r><w:rPr><w:strike w:val="false"/></w:rPr><w:t>kept-false</w:t></w:r>
          <w:r><w:rPr><w:strike w:val="0"/></w:rPr><w:t> kept-zero</w:t></w:r>
          <w:r><w:rPr><w:strike w:val="off"/></w:rPr><w:t> kept-off</w:t></w:r>
          <w:r><w:rPr><w:strike/></w:rPr><w:t> GONE</w:t></w:r>
          <w:r><w:rPr><w:strike w:val="true"/></w:rPr><w:t> ALSO-GONE</w:t></w:r>
        </w:p>""".encode("utf-8")
    )

    paras = _read_paragraph(paragraph, resolver, None, 0, drop_strikethrough=True)

    assert len(paras) == 1
    assert paras[0].text == "kept-false kept-zero kept-off"
    assert "GONE" not in paras[0].text


# --------------------------------------------------------------------------
# Soft breaks (Q2)
# --------------------------------------------------------------------------


@pytest.mark.parametrize(
    "sample, source_paras, expected_blocks",
    [(PN_CST_38, 85, 92), (CARNE_LEAO, 109, 112), (PAR_COSIT, 100, 104)],
)
def test_soft_break_splits_paragraph(
    samples_dir, sample, source_paras, expected_blocks
):
    """A `<w:br/>` inside a paragraph starts a new block.

    `pn_cst_38` has 85 source `<w:p>` elements and 7 soft breaks, yielding 92
    blocks. Leaving them joined would glue a heading to the text beneath it.
    Source counts are recomputed from the OOXML rather than hard-coded twice.
    """
    from docx import Document
    from docx.oxml.ns import qn

    document = Document(str(samples_dir / sample))
    counted = sum(1 for child in document.element.body if child.tag == qn("w:p"))
    assert counted == source_paras, "sample changed: source paragraph count moved"

    doc = read_docx(samples_dir / sample)
    assert len(doc.blocks) == expected_blocks
    assert len(doc.blocks) > counted, "no split occurred"


# --------------------------------------------------------------------------
# Hyperlinks (Q2)
# --------------------------------------------------------------------------


def test_hyperlink_href_captured(carne_leao):
    """11 inlines carry a resolved `href`; the rest carry `None`."""
    linked = [i for p in carne_leao.paragraphs for i in p.inlines if i.href]

    assert len(linked) == 11
    assert all(isinstance(i.href, str) and i.href for i in linked)
    assert all(not i.is_plain for i in linked)
    assert any(i.href.startswith("http") for i in linked), (
        f"expected external targets, got {sorted({i.href for i in linked})}"
    )


# --------------------------------------------------------------------------
# Tables and block order (§2.4)
# --------------------------------------------------------------------------


@pytest.mark.parametrize("sample, index, shape", TABLE_SAMPLES)
def test_tables_extracted(samples_dir, sample, index, shape):
    """Row/cell shape survives, and cells hold `StyledPara`s, not raw strings."""
    doc = read_docx(samples_dir / sample)

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.shape == shape
    assert len(table.rows) == shape[0]
    assert all(
        isinstance(p, StyledPara)
        for row in table.rows
        for cell in row.cells
        for p in cell.paras
    )
    assert any(cell.text for row in table.rows for cell in row.cells)


@pytest.mark.parametrize("sample, index, shape", TABLE_SAMPLES)
def test_table_block_order_preserved(samples_dir, sample, index, shape):
    """Tables are interleaved mid-document, never appended.

    `python-docx`'s `.paragraphs`/`.tables` accessors flatten this away, so the
    reader walks `body` children directly. `index` is assigned *after*
    soft-break splitting, so it addresses the blocks that actually exist.
    """
    doc = read_docx(samples_dir / sample)
    table = doc.tables[0]

    assert table.index == index
    assert isinstance(doc.blocks[index], StyledTable)
    assert doc.blocks[index] is table
    assert 0 < index < len(doc.blocks) - 1, "table is neither first nor last"
    assert [b.index for b in doc.blocks] == list(range(len(doc.blocks)))


@pytest.mark.parametrize(
    "sample",
    [s for s in ALL_SAMPLES if s not in {RESP, SUMULA_STJ, PAR_COSIT}],
)
def test_samples_without_tables_report_none(samples_dir, sample):
    """The complement of `TABLE_SAMPLES`: twelve documents with no table."""
    doc = read_docx(samples_dir / sample)

    assert doc.tables == ()
    assert all(isinstance(b, StyledPara) for b in doc.blocks)


# --------------------------------------------------------------------------
# Alignment
# --------------------------------------------------------------------------


def test_alignment_captured(parecer_93):
    """`w:jc/@w:val` survives on the paragraphs that declare it.

    NOTE — spec §6 predicted `alignment == "both"` (justified) here. The real
    document declares no `w:jc w:val="both"` at all: its explicit alignments are
    `center`, `left` and `right`, with the majority inheriting from the style
    and reporting `None`. Asserted as measured.
    """
    values = {p.alignment for p in parecer_93.paragraphs}

    assert values == {None, "center", "left", "right"}
    aligned = [p for p in parecer_93.paragraphs if p.alignment is not None]
    assert aligned
    assert all(isinstance(p.alignment, str) for p in aligned)


def test_alignment_none_when_undeclared(samples_dir):
    """A document with no `w:jc` anywhere reports `None` throughout — absence is
    represented as absence, not as a fabricated default."""
    doc = read_docx(samples_dir / "sumula_carf_42.docx")

    assert {p.alignment for p in doc.paragraphs} == {None}


# --------------------------------------------------------------------------
# Corpus-wide invariants
# --------------------------------------------------------------------------


@pytest.mark.parametrize("sample", ALL_SAMPLES)
def test_all_samples_read(samples_dir, sample):
    """Exit criterion: all 15 samples ingest without error, each ≥1 block."""
    doc = read_docx(samples_dir / sample)

    assert len(doc.blocks) >= 1
    assert doc.paragraphs, f"{sample} produced no paragraph"
    assert all(isinstance(b, (StyledPara, StyledTable)) for b in doc.blocks)
    assert [b.index for b in doc.blocks] == list(range(len(doc.blocks)))


def test_sample_inventory_is_complete(samples_dir):
    """The parametrisation above covers the corpus — no sample skipped."""
    on_disk = sorted(p.name for p in samples_dir.glob("*.docx"))

    assert on_disk == sorted(ALL_SAMPLES)
    assert len(on_disk) == 15


def test_empty_paragraphs_retained(parecer_93):
    """Blank paragraphs are kept, flagged by `is_empty`.

    Cycle 3 may read blank lines as front/back-matter separators; dropping them
    at ingestion is irreversible (spec §3, non-blocking decisions).
    """
    empties = [p for p in parecer_93.paragraphs if p.is_empty]

    assert empties
    assert all(p.text.strip() == "" for p in empties)
    # They still carry their structural evidence, not just emptiness.
    assert all(isinstance(p.indent_effective, int) for p in empties)


@pytest.mark.parametrize("sample", ALL_SAMPLES)
def test_source_is_bare_filename(samples_dir, sample):
    """`source` must never encode a checkout path, or goldens stop being
    machine-independent."""
    doc = read_docx(samples_dir / sample)

    assert doc.source == sample
    assert "/" not in doc.source
    assert "\\" not in doc.source
    assert not Path(doc.source).is_absolute()


# --------------------------------------------------------------------------
# Error handling
# --------------------------------------------------------------------------


def test_missing_file_raises(tmp_path):
    with pytest.raises(DocxReadError) as excinfo:
        read_docx(tmp_path / "absent.docx")

    assert "absent.docx" in str(excinfo.value)


def test_non_docx_file_raises(tmp_path):
    """A `.txt` renamed or not, the reader must fail loudly rather than emit an
    empty `StyledDoc` that looks like a successfully parsed blank document."""
    bogus = tmp_path / "not_a_document.txt"
    bogus.write_text("Isto não é um DOCX.\n", encoding="utf-8")

    with pytest.raises(DocxReadError):
        read_docx(bogus)


def test_directory_raises(tmp_path):
    with pytest.raises(DocxReadError):
        read_docx(tmp_path)
