"""Degenerate ``.docx`` documents, built with ``python-docx`` at test time.

Plan §9.1's robustness layer asks what the parser does with documents that
carry *no* usable structure — and with files that are not documents at all.
Neither question can be put to the corpus: all fifteen samples are real,
well-formed legal texts with at least some structure, which is precisely why
they were chosen. So the degenerate cases are constructed here, following the
A-1.3 / A-4.6 precedent for constructs the corpus cannot reach.

**Why built rather than committed.** Reconciliation answer R-11 settled this.
A committed ``.docx`` is a ZIP of XML: a reviewer cannot see from the diff what
a case *is*, and a case that stops exercising what it was written for — because
the reader learned to look somewhere else — fails silently as a passing test.
Here each case's construction is three readable lines, and the thing under test
is visible in them.

**What "degenerate" means here.** Not *malformed* — the ten
:data:`DEGENERATE_CASES` are all perfectly valid DOCX files that Word would
open without complaint. They are degenerate in the mathematical sense: each is
a limiting case where some signal the pipeline reads is absent, uniform, or
maximal. A document with no headings gives the hierarchy nothing to infer from;
one with nothing but headings gives it structure with no prose to hang on it;
one with four identical headings tests that identical input still yields
distinct ids (invariant #5). The three *malformed* files —
:func:`corrupt_docx`, :func:`truncated_docx`, :func:`zip_that_is_not_a_docx` —
are separate, because they exercise a different contract: a clean
``DocxReadError`` rather than a graceful rendering.

Each builder documents, in its own docstring, what its case is a degenerate
case *of* — the pipeline stage it starves and the invariant that must survive
the starvation.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

__all__ = [
    "DEGENERATE_CASES",
    "build_all",
    "build_case",
    "corrupt_docx",
    "truncated_docx",
    "zip_that_is_not_a_docx",
]

#: Every degenerate case, by name. The name is what a parametrised test ids by,
#: so it appears verbatim in failure output — hence the flat, descriptive names
#: rather than numbers. Order is the order they are introduced below: from the
#: emptiest document upwards.
DEGENERATE_CASES: tuple[str, ...] = (
    "empty",
    "single_paragraph",
    "headings_only",
    "no_headings",
    "deeply_nested_lists",
    "unlabelled_prose",
    "tables_only",
    "whitespace_only",
    "one_giant_paragraph",
    "duplicate_headings",
)

#: ``one_giant_paragraph``'s size. Several thousand characters, per the spec —
#: large enough that a per-character or per-token cost in any stage would show
#: up against the 21-second suite baseline, small enough that building ten
#: documents stays cheap. The sentence is Portuguese legal boilerplate so the
#: profile scorers see plausible input rather than lorem ipsum.
_GIANT_SENTENCE = (
    "A autoridade fiscal competente devera observar o disposto neste ato "
    "declaratorio interpretativo para fins de apuracao do tributo devido, "
    "ressalvadas as hipoteses expressamente excepcionadas na legislacao "
    "vigente. "
)
_GIANT_REPEATS = 40  # ≈ 9 500 characters


def _new_document():
    """A blank ``python-docx`` document.

    ``Document()`` with no argument opens the library's default template, whose
    body holds a ``<w:sectPr>`` and nothing else — no paragraphs. That is what
    makes the ``empty`` case reachable at all: there is no way to *delete* the
    last paragraph of a Word document, so an empty document has to be one that
    never had one.
    """
    return Document()


# ---------------------------------------------------------------------------
# The ten degenerate documents
# ---------------------------------------------------------------------------


def _build_empty(path: Path) -> None:
    """A document with no paragraphs at all.

    Degenerate case *of ingestion*: ``read_docx`` returns a ``StyledDoc`` whose
    ``blocks`` is empty, so every later stage receives nothing to work with.
    The whole pipeline must still produce a document — LexML requires at least
    a ``Metadado`` and a ``PartePrincipal``, and neither can be derived from
    text that does not exist. This is the case that proves the fallbacks are
    fallbacks and not merely lightly-exercised branches.
    """
    _new_document().save(str(path))


def _build_single_paragraph(path: Path) -> None:
    """Exactly one paragraph of prose.

    Degenerate case *of hierarchy inference*: one block cannot establish a
    pattern, so every signal Cycle 4 fuses — style census, indent modality,
    label sequences — has a sample size of one. The tree must come back flat
    and the paragraph must survive into the output intact (invariant #2).
    """
    document = _new_document()
    document.add_paragraph(
        "Trata-se de consulta formulada acerca da incidencia do imposto."
    )
    document.save(str(path))


def _build_headings_only(path: Path) -> None:
    """Three ``Heading 1`` paragraphs, no body text under any of them.

    Degenerate case *of the section-with-content assumption*: structure exists
    and is unambiguous — outline levels say so — but no section has prose. An
    emitter that assumes every ``Agrupamento`` has at least one child ``<p>``
    breaks here, and LexML's content models are strict enough that an empty
    grouping may not be emittable at all. Whatever the answer, it must be
    reached by construction rather than by crashing.
    """
    document = _new_document()
    for title in ("RELATORIO", "FUNDAMENTACAO", "CONCLUSAO"):
        document.add_paragraph(title, style="Heading 1")
    document.save(str(path))


def _build_no_headings(path: Path) -> None:
    """Six ``Normal`` paragraphs — no heading style anywhere, and no labels.

    Degenerate case *of every structural signal at once*: no outline level, no
    numbering, no label ("Art. 1º", "I -", "1."), no indent variation, one
    style throughout. This is the shape the flat fallback exists for, and the
    single most likely shape among 300+ unseen documents, since a great many
    older scanned-and-retyped acts carry no styling whatsoever. The hierarchy
    must report flat rather than invent a grouping (invariant #8).
    """
    document = _new_document()
    for n in range(6):
        document.add_paragraph(
            f"Paragrafo corrido numero {n} sem qualquer marcacao estrutural, "
            "escrito de forma continua como em um documento datilografado."
        )
    document.save(str(path))


def _build_deeply_nested_lists(path: Path) -> None:
    """A single Word numbered list six levels deep.

    Degenerate case *of ``ilvl`` nesting* (amendment A-4.6), pushed past the
    depth any sample reaches: ``CARNE_LEAO``'s ``ilvl=1`` and ``ilvl=2``
    paragraphs are eleven blocks apart and so are two lists, not one nested
    one. Here all six items share one ``w:numId``, which is what makes Word —
    and ``hierarchy/tree.py`` — treat them as one list rather than six.

    Six levels is deliberately deeper than LexML's own hierarchy: it forces the
    question of what happens when the source nests further than the target
    schema can express, which a flat emitter must answer by degrading rather
    than by dropping items.

    **The numbering is written onto the XML directly.** ``python-docx`` has no
    public API for paragraph-level ``w:numPr`` — ``add_paragraph(style="List
    Number")`` applies the *style*, whose numbering lives in ``numbering.xml``
    and which therefore carries no per-paragraph ``w:ilvl`` at all. Since the
    reader (``docx_reader._read_paragraph``) reads ``w:numPr/w:numId`` and
    ``w:numPr/w:ilvl`` off the paragraph's own ``pPr``, the style alone would
    leave ``num_id`` and ``ilvl`` both ``None`` and this fixture would test
    nothing. The elements are added through ``get_or_add_numPr`` /
    ``get_or_add_ilvl`` / ``get_or_add_numId``, which are ``python-docx``'s
    schema-ordered element accessors, so ``w:numPr`` lands in its correct
    position within ``w:pPr`` and the file stays valid OOXML.

    ``numId`` 3 is the default template's own decimal list definition; using an
    existing definition rather than inventing one keeps the document openable
    in Word, which matters because a reviewer's first move on a suspicious
    fixture is to open it.
    """
    document = _new_document()
    for level in range(6):
        paragraph = document.add_paragraph(
            f"Item no nivel {level} da lista aninhada.", style="List Number"
        )
        numPr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        numPr.get_or_add_ilvl().set(qn("w:val"), str(level))
        numPr.get_or_add_numId().set(qn("w:val"), "3")
    document.save(str(path))


def _build_unlabelled_prose(path: Path) -> None:
    """Paragraphs that *look* structural but carry no label or numbering.

    Degenerate case *of typography-as-evidence*: bold, centred, all-caps lines
    are exactly the shape Cycle 4's typography signal reads as a heading, but
    here they carry no number, no ordinal, no "Art.", nothing a label parser
    can anchor to. The document is a trap for a rule that promotes on
    appearance alone, and it is not a contrived one — untitled all-caps
    section markers are common in older *pareceres*.

    Body paragraphs are interleaved so the caps lines have something to be
    headings *of*, which is what makes a wrong promotion produce visibly wrong
    output rather than a harmless no-op.
    """
    document = _new_document()
    for marker in ("DA COMPETENCIA", "DO MERITO", "DA CONCLUSAO"):
        heading = document.add_paragraph()
        heading.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER, without the import
        run = heading.add_run(marker)
        run.bold = True
        document.add_paragraph(
            "Texto corrido subordinado ao marcador anterior, sem numeracao "
            "nem rotulo de qualquer especie."
        )
    document.save(str(path))


def _build_tables_only(path: Path) -> None:
    """One 3x3 table and nothing else.

    Degenerate case *of the paragraph assumption*: ``StyledDoc.text`` excludes
    table text by design (cell text is reached through ``StyledCell.text``), so
    to any stage that reads ``.text`` this document is empty while plainly
    carrying content. Text conservation (invariant #2) is the sharp edge —
    whatever the emitter does with the table, the nine cell strings must not
    vanish and must not be duplicated.

    ``python-docx``'s ``add_table`` seeds each cell with one empty paragraph,
    so writing to ``cell.text`` fills that paragraph rather than adding a
    second one; the cells hold exactly one paragraph each.
    """
    document = _new_document()
    table = document.add_table(rows=3, cols=3)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.text = f"celula {row_index}-{column_index}"
    document.save(str(path))


def _build_whitespace_only(path: Path) -> None:
    """Paragraphs containing only spaces, tabs and non-breaking spaces.

    Degenerate case *of normalisation*: ``normalize_text`` collapses every one
    of these to a single space or to nothing, so the document has blocks but no
    text — a state distinct from both ``empty`` (no blocks) and
    ``single_paragraph`` (blocks with text). ``StyledPara.is_empty`` is true for
    every block here, which is exactly the predicate segmentation uses to find
    front/back-matter separators; a document that is *nothing but* separators
    is the boundary condition that predicate was never asked about.

    NBSP (U+00A0) is included because it is the one whitespace character the
    reference parser names explicitly (Java's ``\\s`` is ASCII-only) and the one
    a copy-paste from a PDF most often introduces.
    """
    document = _new_document()
    for blank in ("   ", "\t\t", "   ", " \t   ", ""):
        document.add_paragraph(blank)
    document.save(str(path))


def _build_one_giant_paragraph(path: Path) -> None:
    """A single paragraph of several thousand characters.

    Degenerate case *of block granularity*: real documents distribute text
    across blocks, and anything quadratic in paragraph length — a repeated
    regex scan, a per-character normalisation pass rebuilt per stage — is
    invisible on fifteen samples of ordinary paragraphs and painful on a
    document that arrived from a bad PDF extraction as one unbroken run. It
    also gives the segment writers a single citable unit far larger than any
    they were sized for.
    """
    document = _new_document()
    document.add_paragraph(_GIANT_SENTENCE * _GIANT_REPEATS)
    document.save(str(path))


def _build_duplicate_headings(path: Path) -> None:
    """The same heading text four times, each with body text beneath it.

    Degenerate case *of id derivation* (invariant #5, id uniqueness). Any
    scheme that derives an id from the heading's text — a slug, a hash, a
    normalised label — collides here on the first duplicate, and a colliding
    ``id`` is a schema violation, not merely untidy output: LexML's ``id`` is
    typed ``xs:ID``. Four repetitions rather than two, so a scheme that
    disambiguates only the *second* occurrence still fails.

    The body text differs per section so a de-duplicating bug that silently
    merged two sections would show as lost text (invariant #2) as well as a
    wrong section count.
    """
    document = _new_document()
    for n in range(4):
        document.add_paragraph("DISPOSICOES GERAIS", style="Heading 1")
        document.add_paragraph(
            f"Conteudo distinto da secao {n}, repetida sob titulo identico."
        )
    document.save(str(path))


#: Name → builder. The dispatch table *is* the case list's definition of truth:
#: :func:`build_case` rejects a name that has no builder, and a builder with no
#: name in :data:`DEGENERATE_CASES` is caught by the assertion below, so the two
#: cannot drift apart unnoticed.
_BUILDERS = {
    "empty": _build_empty,
    "single_paragraph": _build_single_paragraph,
    "headings_only": _build_headings_only,
    "no_headings": _build_no_headings,
    "deeply_nested_lists": _build_deeply_nested_lists,
    "unlabelled_prose": _build_unlabelled_prose,
    "tables_only": _build_tables_only,
    "whitespace_only": _build_whitespace_only,
    "one_giant_paragraph": _build_one_giant_paragraph,
    "duplicate_headings": _build_duplicate_headings,
}

assert tuple(_BUILDERS) == DEGENERATE_CASES, "case list and builders disagree"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_case(name: str, directory: Path) -> Path:
    """Write the named degenerate document into ``directory``; return its path.

    The file is named ``<name>.docx``, so a failure report that names a
    temporary path still says which case failed.

    Raises:
        KeyError: ``name`` is not one of :data:`DEGENERATE_CASES`. Raised with
            the known names in the message, because the usual cause is a typo
            in a parametrised test id and the fix is to see the list.
    """
    try:
        builder = _BUILDERS[name]
    except KeyError:
        raise KeyError(
            f"unknown degenerate case {name!r}; known cases: "
            + ", ".join(DEGENERATE_CASES)
        ) from None
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"{name}.docx"
    builder(path)
    return path


def build_all(directory: Path) -> dict[str, Path]:
    """Every case, written once. Returns ``name -> path``.

    Cheap enough to call from a session-scoped fixture: ten documents of a few
    kilobytes each, built from the same in-memory template. The point of
    building them once is not the write cost but the read cost — the robustness
    suite runs nine stages over each of the ten, and rebuilding per test would
    multiply that by nine for no gain, since nothing downstream mutates the
    files.

    Insertion order follows :data:`DEGENERATE_CASES`, so iterating the result
    gives deterministic test ordering (invariant #4).
    """
    directory = Path(directory)
    return {name: build_case(name, directory) for name in DEGENERATE_CASES}


# ---------------------------------------------------------------------------
# Malformed files — not documents at all
# ---------------------------------------------------------------------------
#
# These three are graded by how far into the read they get before failing, and
# each must fail as a `DocxReadError` rather than as whatever `zipfile` or
# `lxml` happens to raise. The distinction matters for the CLI: §3.5's exit
# codes give "unreadable source" a code of its own, and a traceback escaping to
# stderr is an E-4 failure, not a cosmetic one.


def corrupt_docx(directory: Path) -> Path:
    """A file of arbitrary bytes with a ``.docx`` suffix.

    Fails at the outermost layer: not a ZIP at all, so the archive's central
    directory is never found. The bytes are fixed rather than randomly
    generated — determinism (invariant #4) applies to fixtures too, and a
    fixture that fails one run in a thousand because it accidentally produced a
    valid ZIP header is worse than useless. The leading ``PK`` is deliberate:
    it makes the file *look* like a ZIP to anything sniffing magic bytes, so
    the error has to come from an actual parse rather than from a guess.
    """
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / "corrupt.docx"
    path.write_bytes(b"PK\x03\x04" + bytes(range(256)) * 8)
    return path


def truncated_docx(directory: Path) -> Path:
    """A real DOCX cut in half.

    Fails one layer in: the file begins as a valid ZIP — correct local file
    headers, readable compressed data — and only the central directory at the
    tail is missing. This is the shape a real interrupted download or a
    truncated email attachment takes, and it is the case a "does it start with
    PK?" check would wave through.

    Built from :func:`_build_one_giant_paragraph` rather than an empty
    document, because a minimal DOCX is small enough that half of it might
    still contain no complete part; a document with several kilobytes of text
    guarantees the truncation lands mid-archive.
    """
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    whole = directory / "_whole_for_truncation.docx"
    _build_one_giant_paragraph(whole)
    data = whole.read_bytes()
    path = directory / "truncated.docx"
    path.write_bytes(data[: len(data) // 2])
    return path


def zip_that_is_not_a_docx(directory: Path) -> Path:
    """A valid ZIP archive with no ``word/document.xml``.

    Fails at the innermost layer: the archive opens, its entries list, and only
    the OPC package structure is wrong. A ``.docx`` is a ZIP, so this is what
    every non-DOCX archive renamed to ``.docx`` looks like — a ``.odt``, a
    ``.pptx``, a plain ``.zip`` of the source text someone meant to convert.

    The entries are plausible-but-wrong rather than nonsense: a
    ``[Content_Types].xml`` and a ``_rels`` part are present, so a check for
    "does it look like OPC?" passes and only the missing main document part
    gives it away.
    """
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / "not_a_docx.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types '
            'xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0"?><Relationships '
            'xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>',
        )
        archive.writestr("readme.txt", "this archive carries no word/document.xml")
    return path
