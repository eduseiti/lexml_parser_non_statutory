"""Ingestion is lossless, deterministic and byte-stable across the 15 samples.

Three invariants are guarded here, in increasing order of importance:

1. **Byte-stability** — ``read_docx(sample).to_json()`` still equals the golden
   committed under ``tests/golden/styled/``. This does not say the output is
   *correct*; it says a change to the reader is visible in a reviewable diff
   instead of arriving silently.
2. **Determinism** — plan invariant #4. The same file read twice gives the same
   bytes, so a golden diff means the code changed, never that a ``set``
   iterated differently today.
3. **Conservation** — :func:`test_text_conservation`, the objective meaning of
   Cycle 1's exit criterion *"ingests losslessly"*. Unlike (1) and (2), this one
   is independent of the goldens: it re-derives the expected text from the
   OOXML with ``python-docx`` and compares. A golden can be wrong and stay
   green forever if it was regenerated from a broken reader; conservation
   cannot, because nothing it checks against is produced by our reader.

Goldens are regenerated only by ``python3 scripts/regen_goldens.py`` — never as
a side effect of running this suite (plan §9.4).
"""

from __future__ import annotations

import json
from collections import Counter
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from lexml_nonstat.ingest import StyledDoc, normalize_text, read_docx

REPO_ROOT = Path(__file__).resolve().parents[2]
SAMPLES_DIR = REPO_ROOT / "samples"
GOLDEN_DIR = REPO_ROOT / "tests" / "golden" / "styled"

#: Sample stems, sorted so test ids are stable and readable
#: (``test_golden_matches[parecer_93_2018_decor_cgu_agu]``).
SAMPLE_STEMS = sorted(p.stem for p in SAMPLES_DIR.glob("*.docx"))

REGEN_HINT = (
    "If this change is intended, run `python3 scripts/regen_goldens.py` and "
    "review the resulting diff before committing — a golden change is a "
    "behaviour change."
)

#: Word toggle semantics: present means on unless explicitly switched off.
_TOGGLE_OFF = frozenset({"false", "0", "off"})


# --------------------------------------------------------------------------
# Caching
#
# 15 samples x 8 tests would re-parse `parecer_93` (450 paragraphs) dozens of
# times. Everything below reads through these module-scoped caches, keeping the
# whole file to a couple of seconds.
# --------------------------------------------------------------------------

_DOC_CACHE: dict[str, StyledDoc] = {}
_GOLDEN_CACHE: dict[str, str] = {}


def sample_path(stem: str) -> Path:
    return SAMPLES_DIR / f"{stem}.docx"


def golden_path(stem: str) -> Path:
    return GOLDEN_DIR / f"{stem}.json"


def read_sample(stem: str) -> StyledDoc:
    """The ingested document, parsed once per session."""
    if stem not in _DOC_CACHE:
        _DOC_CACHE[stem] = read_docx(sample_path(stem))
    return _DOC_CACHE[stem]


def read_golden(stem: str) -> str:
    """The committed golden, verbatim — no parsing, no normalisation."""
    if stem not in _GOLDEN_CACHE:
        _GOLDEN_CACHE[stem] = golden_path(stem).read_text(encoding="utf-8")
    return _GOLDEN_CACHE[stem]


def test_samples_are_present():
    """Guards the guard: a parametrisation over an empty glob is 0 silent passes."""
    assert len(SAMPLE_STEMS) == 15, f"expected 15 samples, found {SAMPLE_STEMS}"


# --------------------------------------------------------------------------
# 1. Golden files
# --------------------------------------------------------------------------


def test_golden_exists_for_every_sample():
    """Set equality in both directions — a missing golden and an orphan golden
    are different bugs, and neither may hide behind a count comparison."""
    samples = set(SAMPLE_STEMS)
    goldens = {p.stem for p in GOLDEN_DIR.glob("*.json")}

    assert samples == goldens, (
        f"samples without a golden: {sorted(samples - goldens)}; "
        f"goldens without a sample: {sorted(goldens - samples)}. {REGEN_HINT}"
    )


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_golden_matches(stem: str):
    """Byte-identical, not merely equivalent.

    Comparing parsed JSON would tolerate reordering, indentation drift and
    ``ensure_ascii`` flips — precisely the changes that make a golden diff
    unreadable for the next reviewer.
    """
    produced = read_sample(stem).to_json()
    committed = read_golden(stem)

    if produced == committed:
        return

    # Point at the first divergence rather than dumping two 450-block files.
    produced_lines = produced.splitlines()
    committed_lines = committed.splitlines()
    detail = "output and golden are identical in every line but differ in length"
    for lineno, (got, want) in enumerate(zip(produced_lines, committed_lines), 1):
        if got != want:
            detail = (
                f"first difference at line {lineno}:\n"
                f"  golden:   {want!r}\n"
                f"  produced: {got!r}"
            )
            break

    pytest.fail(
        f"{stem}: read_docx output does not match "
        f"tests/golden/styled/{stem}.json "
        f"({len(committed_lines)} golden lines vs {len(produced_lines)} produced).\n"
        f"{detail}\n{REGEN_HINT}"
    )


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_goldens_are_valid_json(stem: str):
    """A golden must survive the round trip back into the model, or it is a
    text file that merely looks like a serialised document."""
    data = json.loads(read_golden(stem))

    assert isinstance(data, dict)
    assert isinstance(data.get("blocks"), list)
    assert StyledDoc.from_dict(data) == read_sample(stem)


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_golden_source_is_bare_filename(stem: str):
    """``source`` is a filename, never a path.

    Goldens are committed. If ``source`` held an absolute path, every developer
    with a different checkout directory would see all 15 goldens as modified,
    and the regeneration policy would collapse into noise.
    """
    source = json.loads(read_golden(stem)).get("source")

    assert source == f"{stem}.docx"
    assert "/" not in source and "\\" not in source, (
        f"golden encodes a path, not a filename: {source!r}"
    )


# --------------------------------------------------------------------------
# 2. Model round-trip and determinism
# --------------------------------------------------------------------------


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_roundtrip_from_dict(stem: str):
    """``from_dict(to_dict(doc)) == doc``, structurally.

    The dataclasses are frozen, so ``==`` compares every field of every inline
    of every paragraph — this is a real equality, not identity. It matters
    because ``to_dict`` omits default-valued fields for readability, and an
    omission that ``from_dict`` cannot reconstruct would be silent data loss
    in the goldens themselves.
    """
    doc = read_sample(stem)

    assert StyledDoc.from_dict(doc.to_dict()) == doc
    assert StyledDoc.from_json(doc.to_json()) == doc


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_determinism(stem: str):
    """Plan invariant #4: two independent reads produce identical JSON.

    Deliberately does *not* use the cached document — the point is to exercise
    the reader twice. Non-determinism here (dict/set iteration order leaking
    into output) would make golden diffs meaningless.
    """
    first = read_docx(sample_path(stem)).to_json()
    second = read_docx(sample_path(stem)).to_json()

    assert first == second, f"{stem}: two reads of the same file disagree"


# --------------------------------------------------------------------------
# 3. Conservation — the independent cross-check
# --------------------------------------------------------------------------


def _toggle_on(element) -> bool:
    """OOXML toggle property: present means on unless ``w:val`` disables it."""
    if element is None:
        return False
    return (element.get(qn("w:val")) or "true").lower() not in _TOGGLE_OFF


def _run_is_struck(run) -> bool:
    rPr = run.find(qn("w:rPr"))
    if rPr is None:
        return False
    return _toggle_on(rPr.find(qn("w:strike"))) or _toggle_on(
        rPr.find(qn("w:dstrike"))
    )


#: Sentinel marking a soft break inside a run's text stream.
_SOFT_BREAK = object()


def _run_pieces(run) -> list:
    """One run as a list of text fragments and ``_SOFT_BREAK`` sentinels.

    Written against the OOXML directly rather than through ``Run.text``, which
    silently drops ``<w:tab/>``, ``<w:br/>`` and ``<w:noBreakHyphen/>`` — the
    very elements the conservation question is about.
    """
    pieces: list = []
    for child in run:
        tag = child.tag
        if tag in (qn("w:t"), qn("w:delText")):
            pieces.append(child.text or "")
        elif tag == qn("w:tab"):
            pieces.append(" ")
        elif tag == qn("w:br"):
            # Only a text-wrapping break splits a paragraph; a page or column
            # break is layout, and its text flows on.
            if (child.get(qn("w:type")) or "textWrapping") == "textWrapping":
                pieces.append(_SOFT_BREAK)
            else:
                pieces.append(" ")
        elif tag == qn("w:cr"):
            pieces.append(_SOFT_BREAK)
        elif tag == qn("w:noBreakHyphen"):
            pieces.append("-")
        elif tag == qn("w:sym"):
            char = child.get(qn("w:char"))
            if char:
                try:
                    pieces.append(chr(int(char, 16)))
                except ValueError:
                    pass
    return pieces


def _source_runs(paragraph):
    """Every ``<w:r>`` that contributes text, in document order.

    Runs nest inside hyperlinks, tracked insertions, smart tags and content
    controls; ``<w:p>``'s direct children are not the whole story.
    """
    for child in paragraph:
        if child.tag == qn("w:r"):
            yield child
        elif child.tag in (
            qn("w:hyperlink"),
            qn("w:ins"),
            qn("w:smartTag"),
            qn("w:sdt"),
        ):
            yield from child.iter(qn("w:r"))


def _source_lines(paragraph, *, drop_struck: bool = True) -> list[str]:
    """One source paragraph as the lines a lossless reader should emit.

    Soft breaks split; struck runs are excluded when ``drop_struck``. Each line
    is normalised with the production ``normalize_text`` and edge-stripped,
    matching what ``StyledPara.text`` can express.
    """
    lines: list[list[str]] = [[]]
    for run in _source_runs(paragraph):
        if drop_struck and _run_is_struck(run):
            continue
        for piece in _run_pieces(run):
            if piece is _SOFT_BREAK:
                lines.append([])
            else:
                lines[-1].append(piece)
    return [normalize_text("".join(line)).strip() for line in lines]


def _expected_texts(stem: str, *, drop_struck: bool = True) -> Counter:
    """Non-empty paragraph texts read straight from the OOXML.

    Walks ``document.element.body`` (not ``Document.paragraphs``, which flattens
    away tables and their ordering) and descends into every ``<w:tbl>``, so body
    and table-cell paragraphs are both covered.
    """
    document = Document(str(sample_path(stem)))
    texts: list[str] = []
    for child in document.element.body:
        if child.tag == qn("w:p"):
            texts.extend(_source_lines(child, drop_struck=drop_struck))
        elif child.tag == qn("w:tbl"):
            for paragraph in child.iter(qn("w:p")):
                texts.extend(_source_lines(paragraph, drop_struck=drop_struck))
    return Counter(t for t in texts if t)


def _actual_texts(doc: StyledDoc) -> Counter:
    """Non-empty paragraph texts as ingested, body and table cells alike."""
    texts: list[str] = []
    for block in doc.blocks:
        paras = (
            [block]
            if hasattr(block, "inlines")
            else [
                para
                for row in block.rows
                for cell in row.cells
                for para in cell.paras
            ]
        )
        texts.extend(normalize_text(p.text).strip() for p in paras)
    return Counter(t for t in texts if t)


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_text_conservation(stem: str):
    """No source text is dropped, invented or duplicated by ingestion.

    This is the objective meaning of Cycle 1's exit criterion *"ingests
    losslessly"*, and the most important assertion in this file. "Lossless"
    here means: **the multiset of non-empty, normalised paragraph texts that
    the reader produces is exactly the multiset the source contains** — every
    paragraph present, present the right number of times, with nothing added.

    Three design choices make the test meaningful rather than circular:

    *Independent derivation.* The expected side is rebuilt from the OOXML with
    ``python-docx`` — walking ``document.element.body`` and descending into
    ``<w:tbl>`` — not from ``StyledDoc``. Only ``normalize_text`` is shared, and
    it must be: comparing a normalised string against an unnormalised one would
    fail on whitespace, not on loss. Everything else is a genuine second
    opinion; the reader does not grade its own homework.

    *Multiset, not substring.* Comparison is by ``collections.Counter`` of
    per-paragraph texts. A substring check over the concatenated document would
    pass while a paragraph was silently emitted twice, or while two paragraphs
    were fused into one — both real regressions this formulation catches. The
    failure report is the symmetric difference, so it names the specific text.

    *Two sanctioned exemptions*, and only two. Both are deliberate design
    decisions from spec §3 Q2, both mirror the reference parser, and both are
    modelled on the expected side rather than excused away:

    a. **Soft-break splitting.** A source paragraph containing ``<w:br/>``
       becomes N ``StyledPara``s. Nothing is lost — the text is redistributed —
       so the expected side splits at the same OOXML elements and the
       concatenation invariant is checked separately by
       :func:`test_soft_break_split_conserves_paragraph_text`, which does not
       depend on *where* the reader chose to split. Affects
       ``par_cosit_26`` (3 breaks), ``pn_cst_38`` (7) and ``CARNE_LEAO`` (3).
    b. **Struck-run dropping.** Runs carrying an active ``<w:strike>`` or
       ``<w:dstrike>`` are removed by design, because the reference parser
       removes them and Cycle 6b's round-trip needs a common baseline. This is
       real deletion, so it is principled only because it is *bounded and
       visible*: the expected side excludes exactly the same runs, the toggle
       semantics are re-implemented here rather than imported, ``sumula_stj_125``
       is the only affected sample (18 runs), and
       :func:`test_struck_text_absent_from_default_output` proves the text is
       retrievable via ``drop_strikethrough=False``. Nothing is unrecoverable.

    Anything else that disappears is a bug, and this test is where it surfaces.
    """
    expected = _expected_texts(stem)
    actual = _actual_texts(read_sample(stem))

    if expected == actual:
        return

    missing = expected - actual  # in the source, absent from the output
    duplicated = actual - expected  # in the output, unaccounted for in the source

    def render(counter: Counter, limit: int = 5) -> str:
        if not counter:
            return "    (none)"
        rows = [
            f"    x{count}: {text[:200]!r}"
            for text, count in list(counter.items())[:limit]
        ]
        if len(counter) > limit:
            rows.append(f"    ... and {len(counter) - limit} more")
        return "\n".join(rows)

    pytest.fail(
        f"{stem}: ingestion is not text-conserving "
        f"({sum(missing.values())} paragraph(s) lost, "
        f"{sum(duplicated.values())} unaccounted for).\n"
        f"  In the source but not in StyledDoc (text was DROPPED):\n"
        f"{render(missing)}\n"
        f"  In StyledDoc but not in the source (text was DUPLICATED or "
        f"ALTERED):\n{render(duplicated)}\n"
        "  This is a reader bug, not a golden staleness issue: the expected "
        "side is derived from the OOXML independently of read_docx. Do not "
        "regenerate goldens to make it pass."
    )


@pytest.mark.parametrize("stem", SAMPLE_STEMS)
def test_soft_break_split_conserves_paragraph_text(stem: str):
    """Splitting redistributes text; it never removes any.

    :func:`test_text_conservation` mirrors the reader's split points on the
    expected side, so on its own it could not tell a *correct* split from a
    *consistently wrong* one. This test closes that gap without depending on
    the split at all: it compares the total per-document text, breaks collapsed
    to spaces, so any placement of the split boundaries passes and any lost
    character fails.
    """
    document = Document(str(sample_path(stem)))
    source_words: Counter = Counter()
    for child in document.element.body:
        if child.tag == qn("w:p"):
            paragraphs = [child]
        elif child.tag == qn("w:tbl"):
            paragraphs = list(child.iter(qn("w:p")))
        else:
            continue
        for paragraph in paragraphs:
            source_words.update(" ".join(_source_lines(paragraph)).split())

    ingested_words: Counter = Counter()
    for text in _actual_texts(read_sample(stem)).elements():
        ingested_words.update(text.split())

    assert source_words == ingested_words, (
        f"{stem}: word multiset differs across the split — "
        f"lost {sorted((source_words - ingested_words).items())[:10]}, "
        f"gained {sorted((ingested_words - source_words).items())[:10]}"
    )


def test_struck_text_absent_from_default_output():
    """The struck-run exemption is bounded and reversible.

    ``sumula_stj_125`` is the only sample with struck runs (18 of them: the
    ordinal markers ``ª``/``º`` inside otherwise-live citations such as
    ``(2ª T, 03.08.1994)``). By default they are dropped, matching the
    reference parser; with ``drop_strikethrough=False`` they come back. Losing
    them by *design* is defensible only while both halves of that sentence
    hold, which is what this asserts.
    """
    stem = "sumula_stj_125"
    dropped = _actual_texts(read_docx(sample_path(stem)))
    retained = _actual_texts(
        read_docx(sample_path(stem), drop_strikethrough=False)
    )

    def joined(counter: Counter) -> str:
        return "\n".join(counter.elements())

    assert "2ª T," not in joined(dropped), (
        "struck ordinal survived the default read"
    )
    assert "2ª T," in joined(retained), (
        "drop_strikethrough=False must retain the struck ordinal"
    )
    # The live text around the struck run is untouched either way.
    assert "2 T, 03.08.1994" in joined(dropped)

    # And with the struck runs retained, the output conserves *them* too.
    assert retained == _expected_texts(stem, drop_struck=False)


def test_only_sumula_stj_has_struck_runs():
    """A tripwire on the exemption's scope.

    :func:`test_text_conservation` excludes struck runs from the expected side
    for every sample. That is only harmless while struck runs stay confined to
    one document whose behaviour is separately asserted — the day a new sample
    carries struck text, this fails and forces the exemption to be re-examined
    rather than inherited.
    """
    affected = {}
    for stem in SAMPLE_STEMS:
        document = Document(str(sample_path(stem)))
        count = sum(
            1
            for paragraph in document.element.body.iter(qn("w:p"))
            for run in _source_runs(paragraph)
            if _run_is_struck(run)
        )
        if count:
            affected[stem] = count

    assert affected == {"sumula_stj_125": 18}, (
        f"struck runs found outside the documented scope: {affected}"
    )
